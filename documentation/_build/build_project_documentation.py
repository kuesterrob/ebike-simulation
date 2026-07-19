from __future__ import annotations

from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
from matplotlib.patches import Circle, FancyArrowPatch, FancyBboxPatch, Polygon

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[2]
DOCUMENTATION_DIR = PROJECT_ROOT / "documentation"
FIGURES_DIR = DOCUMENTATION_DIR / "figures"
OUTPUT_FILE = DOCUMENTATION_DIR / "Projektdokumentation_E-Bike-Simulation.docx"
UML_FILE = FIGURES_DIR / "uml_klassendiagramm.png"
ACTIVITY_FILE = FIGURES_DIR / "aktivitaetsdiagramm_simulation.png"


# Design preset: compact_reference_guide.
# Named overrides: A4 page geometry for a German school report, Calibri-based
# editorial cover, and 2.5 cm margins. All tables use the resulting exact A4
# content width rather than the Letter-width base token.
NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MID_BLUE = "5B9BD5"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7FB"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "6B7280"
DARK = "1F2937"
WHITE = "FFFFFF"
GREEN = "2F855A"
ORANGE = "C26D21"
RED = "A33A3A"

CONTENT_WIDTH_DXA = 9072
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_LEFT_RIGHT = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (
        ("top", CELL_MARGIN_TOP_BOTTOM),
        ("bottom", CELL_MARGIN_TOP_BOTTOM),
        ("start", CELL_MARGIN_LEFT_RIGHT),
        ("end", CELL_MARGIN_LEFT_RIGHT),
    ):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "B7C5D2", size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA} DXA")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def style_table(table, widths_dxa: list[int]) -> None:
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    repeat_table_header(table.rows[0])

    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.10
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
                    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = rgb(DARK)
                    if row_index == 0:
                        run.bold = True


def paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_left_border(paragraph, color: str, size: str = "18") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)


def add_field(paragraph, instruction: str, fallback: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.extend([begin, instr, separate])
    if fallback:
        value_run = paragraph.add_run(fallback)
        value_run.font.name = "Calibri"
        value_run.font.size = Pt(9)
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def create_numbering(document: Document, kind: str) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "hybridMultilevel")
    abstract.append(multi)

    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl_text = OxmlElement("w:lvlText")
        if kind == "bullet":
            lvl_text.set(qn("w:val"), ("•", "–", "◦")[level])
        else:
            lvl_text.set(qn("w:val"), f"%{level + 1}.")
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        text_indent = 540 + level * 420
        tab.set(qn("w:pos"), str(text_indent))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(text_indent))
        ind.set(qn("w:hanging"), "270")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Calibri")
        r_fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(r_fonts)
        lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
        abstract.append(lvl)

    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def set_numbering(paragraph, num_id: int, level: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def add_list(document: Document, items: list[str], kind: str = "bullet", level: int = 0) -> None:
    num_id = create_numbering(document, kind)
    for item in items:
        paragraph = document.add_paragraph(style="Normal")
        set_numbering(paragraph, num_id, level)
        paragraph.add_run(item)


def add_rich_list(document: Document, items: list[tuple[str, str]], kind: str = "bullet") -> None:
    num_id = create_numbering(document, kind)
    for label, text in items:
        paragraph = document.add_paragraph(style="Normal")
        set_numbering(paragraph, num_id, 0)
        lead = paragraph.add_run(label)
        lead.bold = True
        paragraph.add_run(text)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.add_run(text)


def add_callout(document: Document, label: str, text: str, color: str = BLUE) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.left_indent = Cm(0.25)
    paragraph.paragraph_format.right_indent = Cm(0.15)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(9)
    paragraph_shading(paragraph, PALE_BLUE)
    paragraph_left_border(paragraph, color)
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = rgb(color)
    paragraph.add_run(text)


def add_code_block(document: Document, lines: str) -> None:
    paragraph = document.add_paragraph(style="Code Block")
    paragraph_shading(paragraph, LIGHT_GRAY)
    for index, line in enumerate(lines.splitlines()):
        if index:
            paragraph.add_run().add_break()
        paragraph.add_run(line)


def add_figure(document: Document, path: Path, caption: str, alt_text: str, width_inches: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(width_inches))
    inline_shape._inline.docPr.set("descr", alt_text)
    inline_shape._inline.docPr.set("title", caption)

    caption_paragraph = document.add_paragraph(style="Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.keep_with_next = False
    caption_paragraph.add_run(caption)


def class_box(ax, x, y, w, h, title, lines, fill="#F4F7FB", edge="#2E74B5"):
    box = FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.006,rounding_size=0.009",
        linewidth=1.4,
        facecolor=fill,
        edgecolor=edge,
        transform=ax.transAxes,
        zorder=2,
    )
    ax.add_patch(box)
    header_y = y + h * 0.68
    ax.plot([x, x + w], [header_y, header_y], color=edge, linewidth=1.0, transform=ax.transAxes, zorder=3)
    ax.text(x + w / 2, y + h * 0.835, title, ha="center", va="center", fontsize=9.5, fontweight="bold", color="#17365D", transform=ax.transAxes, zorder=4)
    body = "\n".join(lines)
    ax.text(x + 0.012, y + h * 0.33, body, ha="left", va="center", fontsize=7.5, color="#1F2937", transform=ax.transAxes, zorder=4)


def arrow(ax, start, end, label="", dashed=True, color="#64748B", curve=0.0):
    patch = FancyArrowPatch(
        start,
        end,
        arrowstyle="-|>",
        mutation_scale=10,
        linewidth=1.1,
        linestyle="--" if dashed else "-",
        color=color,
        connectionstyle=f"arc3,rad={curve}",
        transform=ax.transAxes,
        zorder=1,
    )
    ax.add_patch(patch)
    if label:
        mid_x = (start[0] + end[0]) / 2
        mid_y = (start[1] + end[1]) / 2
        ax.text(mid_x, mid_y + 0.012, label, fontsize=6.8, color=color, ha="center", va="center", transform=ax.transAxes, zorder=5, bbox=dict(facecolor="white", edgecolor="none", pad=0.5))


def inheritance_arrow(ax, start, end):
    ax.plot([start[0], end[0]], [start[1], end[1] - 0.018], color="#17365D", linewidth=1.2, transform=ax.transAxes, zorder=1)
    triangle = Polygon(
        [[end[0], end[1]], [end[0] - 0.009, end[1] - 0.018], [end[0] + 0.009, end[1] - 0.018]],
        closed=True,
        facecolor="white",
        edgecolor="#17365D",
        linewidth=1.2,
        transform=ax.transAxes,
        zorder=3,
    )
    ax.add_patch(triangle)


def create_uml_diagram(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(13.5, 9.2), dpi=190)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    ax.text(0.02, 0.975, "Vereinfachtes UML-Klassendiagramm", fontsize=15, fontweight="bold", color="#17365D", transform=ax.transAxes)
    ax.text(0.02, 0.945, "Kernklassen, Vererbung und wichtigste Abhängigkeiten", fontsize=9.5, color="#64748B", transform=ax.transAxes)

    class_box(ax, 0.03, 0.78, 0.17, 0.12, "main.py", ["+ main()", "+ run_study()", "+ Menüfunktionen"], fill="#FFF7E8", edge="#C26D21")
    class_box(ax, 0.38, 0.76, 0.24, 0.16, "BikeSimulator", ["- gps_file, Parameter", "+ run(): dict", "- _prepare_route_data()", "- _simulate_battery_variant()"], fill="#EAF3FB", edge="#2E74B5")
    class_box(ax, 0.77, 0.78, 0.19, 0.12, "Ausgabemodule", ["plotter.py", "reporting.console", "reporting.pdf_report"], fill="#EDF8F2", edge="#2F855A")

    class_box(ax, 0.03, 0.55, 0.17, 0.13, "GPSReader", ["+ load_file()", "+ calculate_distances()", "+ get_stats()"], fill="#F8FAFC", edge="#64748B")
    class_box(ax, 0.22, 0.55, 0.17, 0.13, "RouteCalculator", ["+ calculate_speed()", "+ calculate_acceleration()", "+ calculate_slope()"], fill="#F8FAFC", edge="#64748B")
    class_box(ax, 0.03, 0.35, 0.17, 0.13, "TripWeather", ["+ get_weather()", "- _fetch()", "Cache + Open-Meteo"], fill="#F8FAFC", edge="#64748B")
    class_box(ax, 0.22, 0.35, 0.17, 0.13, "Reverse_Geocoder", ["+ get_results()", "+ geoapify_bulk()", "Cache + Geoapify"], fill="#F8FAFC", edge="#64748B")
    class_box(ax, 0.03, 0.15, 0.17, 0.13, "MovementDirection", ["+ calculate()", "- _bearing()"], fill="#F8FAFC", edge="#64748B")
    class_box(ax, 0.22, 0.15, 0.17, 0.13, "Cleaner / GPSMap", ["+ clean_places()", "+ save()", "Datenbereinigung + Karte"], fill="#F8FAFC", edge="#64748B")

    class_box(ax, 0.44, 0.55, 0.16, 0.12, "BatteryBase", ["{abstract}", "+ apply_current()", "+ voltage()"], fill="#F3F0FA", edge="#60408F")
    class_box(ax, 0.44, 0.36, 0.16, 0.13, "BatteryPack", ["+ apply_current()", "+ voltage()", "+ update_temperature()", "+ maximum_charge_current()"], fill="#F3F0FA", edge="#60408F")
    class_box(ax, 0.41, 0.14, 0.14, 0.12, "LiPoBatteryPack", ["+ voltage()", "OCV: SoC^0,3"], fill="#F7F3FC", edge="#60408F")
    class_box(ax, 0.57, 0.14, 0.14, 0.12, "NMCBatteryPack", ["+ voltage()", "OCV: SoC^0,834"], fill="#F7F3FC", edge="#60408F")

    class_box(ax, 0.73, 0.56, 0.19, 0.13, "Motor", ["+ calculate()", "Kräfte, Leistung,", "Drehmoment und Strom"], fill="#FFF5F5", edge="#A33A3A")
    class_box(ax, 0.73, 0.35, 0.22, 0.14, "RegenerativeBrakingController", ["+ distribute()", "+ calculate_charge_current()", "Akku → Widerstand → Bremse"], fill="#FFF5F5", edge="#A33A3A")
    class_box(ax, 0.76, 0.14, 0.17, 0.13, "BrakeResistor", ["+ maximum_power()", "+ update_temperature()", "+ dissipated_energy_wh"], fill="#FFF5F5", edge="#A33A3A")

    arrow(ax, (0.20, 0.84), (0.38, 0.84), "startet", dashed=False, color="#C26D21")
    arrow(ax, (0.62, 0.84), (0.77, 0.84), "liefert Ergebnisse", dashed=True, color="#2F855A")

    for start, end in [
        ((0.38, 0.80), (0.20, 0.62)),
        ((0.38, 0.79), (0.39, 0.62)),
        ((0.42, 0.76), (0.20, 0.42)),
        ((0.43, 0.76), (0.39, 0.42)),
        ((0.45, 0.76), (0.20, 0.22)),
        ((0.46, 0.76), (0.39, 0.22)),
    ]:
        arrow(ax, start, end, dashed=True, color="#94A3B8")

    arrow(ax, (0.62, 0.80), (0.73, 0.63), "berechnet", dashed=True, color="#A33A3A")
    arrow(ax, (0.62, 0.77), (0.73, 0.42), "verteilt Bremsleistung", dashed=True, color="#A33A3A")
    arrow(ax, (0.57, 0.76), (0.52, 0.67), "erzeugt", dashed=True, color="#60408F")
    arrow(ax, (0.84, 0.35), (0.84, 0.27), "nutzt", dashed=True, color="#A33A3A")
    arrow(ax, (0.73, 0.39), (0.60, 0.40), "lädt", dashed=True, color="#60408F")

    inheritance_arrow(ax, (0.52, 0.49), (0.52, 0.55))
    inheritance_arrow(ax, (0.48, 0.26), (0.50, 0.36))
    inheritance_arrow(ax, (0.64, 0.26), (0.56, 0.36))

    ax.plot([0.03, 0.97], [0.085, 0.085], color="#D6DEE8", linewidth=0.8, transform=ax.transAxes)
    ax.text(0.03, 0.045, "Durchgezogene Linie: Aufruf | gestrichelte Linie: Abhängigkeit | hohles Dreieck: Vererbung", fontsize=8, color="#64748B", transform=ax.transAxes)

    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def activity_box(ax, x, y, w, h, text, fill="#F4F7FB", edge="#2E74B5", fontsize=8.6):
    box = FancyBboxPatch((x - w / 2, y - h / 2), w, h, boxstyle="round,pad=0.04,rounding_size=0.08", facecolor=fill, edgecolor=edge, linewidth=1.4)
    ax.add_patch(box)
    ax.text(x, y, text, ha="center", va="center", fontsize=fontsize, color="#1F2937", wrap=True)


def activity_diamond(ax, x, y, w, h, text, fill="#FFF7E8", edge="#C26D21", fontsize=8.4):
    points = [(x, y + h / 2), (x + w / 2, y), (x, y - h / 2), (x - w / 2, y)]
    diamond = Polygon(points, closed=True, facecolor=fill, edgecolor=edge, linewidth=1.4)
    ax.add_patch(diamond)
    ax.text(x, y, text, ha="center", va="center", fontsize=fontsize, color="#1F2937", wrap=True)


def activity_arrow(ax, start, end, label="", color="#64748B", connection="arc3,rad=0"):
    arrow_patch = FancyArrowPatch(start, end, arrowstyle="-|>", mutation_scale=11, linewidth=1.25, color=color, connectionstyle=connection)
    ax.add_patch(arrow_patch)
    if label:
        x = (start[0] + end[0]) / 2
        y = (start[1] + end[1]) / 2
        ax.text(x, y + 0.13, label, fontsize=7.3, color=color, ha="center", va="center", bbox=dict(facecolor="white", edgecolor="none", pad=0.8))


def create_activity_diagram(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(8.4, 13.2), dpi=200)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 27)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    ax.text(0.45, 26.45, "Aktivitätsdiagramm der Einzelsimulation", fontsize=14.5, fontweight="bold", color="#17365D")
    ax.text(0.45, 25.95, "Die Parameterstudie wiederholt denselben Ablauf für 20 Parametersätze.", fontsize=8.8, color="#64748B")

    ax.add_patch(Circle((5, 25.15), 0.18, facecolor="#17365D", edgecolor="#17365D"))
    activity_arrow(ax, (5, 24.96), (5, 24.38))
    activity_box(ax, 5, 23.85, 5.8, 0.8, "BikeSimulator erzeugen und Parameter validieren")
    activity_arrow(ax, (5, 23.44), (5, 22.86))
    activity_box(ax, 5, 22.05, 7.4, 1.25, "Routendaten vorbereiten\nCSV lesen · 3D-Distanzen und Zeitdifferenzen · Karte · Orte und Wetter aus Cache/API · Fahrtrichtung · Geschwindigkeit · Beschleunigung · Steigung · Luftdichte", fill="#EEF6FC", fontsize=8.2)
    activity_arrow(ax, (5, 21.42), (5, 20.8))
    activity_box(ax, 5, 20.2, 6.2, 0.9, "Motormodell berechnen\nKräfte → Leistung → Drehmoment → Antriebs- bzw. Bremsbedarf")
    activity_arrow(ax, (5, 19.74), (5, 19.18))
    activity_box(ax, 5, 18.65, 6.3, 0.8, "LiPo, NMC, Rekuperationscontroller und zwei Bremswiderstände erzeugen")
    activity_arrow(ax, (5, 18.24), (5, 17.66))
    activity_box(ax, 5, 17.15, 4.7, 0.72, "Nächste Akkuvariante wählen", fill="#F7F3FC", edge="#60408F")
    activity_arrow(ax, (5, 16.78), (5, 16.18))
    activity_box(ax, 5, 15.7, 4.7, 0.72, "Nächsten Streckenabschnitt lesen", fill="#F7F3FC", edge="#60408F")
    activity_arrow(ax, (5, 15.33), (5, 14.76))
    activity_diamond(ax, 5, 13.92, 4.1, 1.35, "Bremsleistung > 0?")

    activity_arrow(ax, (3.05, 13.92), (2.35, 13.15), "nein")
    activity_arrow(ax, (6.95, 13.92), (7.65, 13.15), "ja")
    activity_box(ax, 2.2, 12.35, 3.7, 1.0, "Antriebsfall\npositiven Motorstrom verwenden", fill="#EDF8F2", edge="#2F855A")
    activity_box(ax, 7.8, 12.35, 3.8, 1.25, "Bremsfall\nLeistung zuerst an Akku, danach an Bremswiderstand und zuletzt an mechanische Bremse verteilen", fill="#FFF5F5", edge="#A33A3A", fontsize=8.0)
    activity_arrow(ax, (2.2, 11.84), (4.25, 10.95), connection="arc3,rad=-0.08")
    activity_arrow(ax, (7.8, 11.71), (5.75, 10.95), connection="arc3,rad=0.08")
    activity_box(ax, 5, 10.55, 6.5, 0.85, "SoC und Akkutemperatur aktualisieren · Bremswiderstand erwärmen/abkühlen")
    activity_arrow(ax, (5, 10.11), (5, 9.52))
    activity_box(ax, 5, 9.05, 5.8, 0.75, "Spannung, Strom, Leistung, Temperaturen und SoC speichern")
    activity_arrow(ax, (5, 8.67), (5, 8.06))
    activity_diamond(ax, 5, 7.35, 4.3, 1.25, "Weitere Streckenabschnitte?")
    activity_arrow(ax, (7.08, 7.35), (9.35, 7.35), "ja")
    activity_arrow(ax, (9.35, 7.35), (9.35, 15.7))
    activity_arrow(ax, (9.35, 15.7), (7.35, 15.7))
    activity_arrow(ax, (5, 6.72), (5, 6.08), "nein")
    activity_diamond(ax, 5, 5.35, 4.2, 1.25, "Weitere Akkuvariante?")
    activity_arrow(ax, (2.9, 5.35), (0.65, 5.35), "ja")
    activity_arrow(ax, (0.65, 5.35), (0.65, 17.15))
    activity_arrow(ax, (0.65, 17.15), (2.65, 17.15))
    activity_arrow(ax, (5, 4.72), (5, 4.14), "nein")
    activity_box(ax, 5, 3.67, 5.8, 0.75, "Kennzahlen berechnen und Ergebnis-Dictionary aufbauen", fill="#EEF6FC")
    activity_arrow(ax, (5, 3.28), (5, 2.7))
    activity_box(ax, 5, 2.18, 6.7, 0.82, "Ausgewählte Terminalwerte ausgeben · optional PDF erzeugen · optional Diagramme anzeigen", fill="#EDF8F2", edge="#2F855A")
    activity_arrow(ax, (5, 1.75), (5, 1.18))
    ax.add_patch(Circle((5, 0.9), 0.24, facecolor="white", edgecolor="#17365D", linewidth=1.5))
    ax.add_patch(Circle((5, 0.9), 0.13, facecolor="#17365D", edgecolor="#17365D"))

    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(DARK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MID_GRAY)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(9)

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style.font.size = Pt(8.7)
    code_style.font.color.rgb = rgb(DARK)
    code_style.paragraph_format.left_indent = Cm(0.35)
    code_style.paragraph_format.right_indent = Cm(0.2)
    code_style.paragraph_format.space_before = Pt(5)
    code_style.paragraph_format.space_after = Pt(8)
    code_style.paragraph_format.line_spacing = 1.05
    code_style.paragraph_format.keep_together = True


def configure_section(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)
    section.different_first_page_header_footer = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("PROJEKTDOKUMENTATION  |  E-BIKE-SIMULATION")
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(8.5)
    run.font.bold = True
    run.font.color.rgb = rgb(MID_GRAY)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Seite ")
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = rgb(MID_GRAY)
    add_field(paragraph, "PAGE", "1")
    run = paragraph.add_run(" von ")
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = rgb(MID_GRAY)
    add_field(paragraph, "NUMPAGES", "1")


def add_cover(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(44)
    paragraph.paragraph_format.space_after = Pt(16)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("ABSCHLUSSPROJEKT PROGRAMMIEREN 1")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = rgb(ORANGE)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(7)
    title_run = title.add_run("Projektdokumentation")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(30)
    title_run.font.bold = True
    title_run.font.color.rgb = rgb(NAVY)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    subtitle_run = subtitle.add_run("GPS-basierte E-Bike-Simulation mit Akkuvergleich und Rekuperation")
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(15)
    subtitle_run.font.color.rgb = rgb(DARK_BLUE)

    description = document.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description.paragraph_format.left_indent = Cm(1.0)
    description.paragraph_format.right_indent = Cm(1.0)
    description.paragraph_format.space_after = Pt(26)
    desc_run = description.add_run(
        "Das Programm verarbeitet GPS- und Umgebungsdaten einer Fahrt, berechnet den mechanischen Leistungsbedarf und simuliert LiPo- und NMC-Akkupacks einschließlich thermischem Verhalten und regenerativem Bremsen."
    )
    desc_run.font.name = "Calibri"
    desc_run.font.size = Pt(11.5)
    desc_run.font.color.rgb = rgb(DARK)

    facts = document.add_paragraph()
    facts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    facts.paragraph_format.space_after = Pt(38)
    fact_run = facts.add_run("2.284 GPS-Punkte   |   94,33 km Referenzfahrt   |   20 bestandene Unit-Tests")
    fact_run.font.name = "Calibri"
    fact_run.font.size = Pt(10.5)
    fact_run.font.bold = True
    fact_run.font.color.rgb = rgb(BLUE)

    for label, value in (
        ("Name", "[Vor- und Nachname ergänzen]"),
        ("Klasse / Kurs", "[Klasse oder Kurs ergänzen]"),
        ("Lehrkraft", "[Name ergänzen]"),
        ("Abgabedatum", "[Datum ergänzen]"),
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(3)
        label_run = paragraph.add_run(f"{label}: ")
        label_run.font.name = "Calibri"
        label_run.font.size = Pt(10.5)
        label_run.font.bold = True
        label_run.font.color.rgb = rgb(NAVY)
        value_run = paragraph.add_run(value)
        value_run.font.name = "Calibri"
        value_run.font.size = Pt(10.5)
        value_run.font.color.rgb = rgb(MID_GRAY)

    document.add_page_break()


def add_toc(document: Document) -> None:
    document.add_heading("Inhaltsverzeichnis", level=1)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    add_field(paragraph, 'TOC \\o "1-2" \\h \\z \\u', "Inhaltsverzeichnis wird beim Öffnen aktualisiert.")
    add_callout(
        document,
        "Hinweis",
        "Die Seitenzahlen im Inhaltsverzeichnis werden in Microsoft Word oder LibreOffice automatisch aktualisiert. Falls nötig: Rechtsklick in das Verzeichnis und 'Feld aktualisieren' wählen.",
        color=ORANGE,
    )
    document.add_page_break()


def add_document_content(document: Document) -> None:
    document.add_heading("1 Einleitung", level=1)
    document.add_heading("1.1 Ziel und Zweck des Projekts", level=2)
    add_body(
        document,
        "Die E-Bike-Simulation ist ein Python-Programm zur Auswertung einer aufgezeichneten Fahrradfahrt. Als Grundlage dient eine CSV-Datei mit geografischen Koordinaten, Höhe, Zeitstempel und Temperatur. Aus diesen Messwerten berechnet das Programm zunächst Strecke, Geschwindigkeit, Beschleunigung, Steigung, Fahrtrichtung und Luftdichte. Anschließend bestimmt ein physikalisches Motormodell die Kräfte, die Antriebsleistung, das Drehmoment sowie den Strombedarf für jeden Streckenabschnitt."
    )
    add_body(
        document,
        "Der zweite Schwerpunkt ist der Vergleich zweier Akkutypen. Ein LiPo- und ein NMC-Akkupack werden mit denselben Fahrdaten belastet. Beide Modelle berücksichtigen Ladezustand, Spannung, Innenwiderstand und Temperatur. In Bremsphasen verteilt ein Rekuperationscontroller die Energie auf Akku, Bremswiderstand und mechanische Bremse. Dadurch lässt sich nicht nur der Energieverbrauch, sondern auch die Rückgewinnung von Bremsenergie untersuchen."
    )
    add_rich_list(
        document,
        [
            ("Datengrundlage: ", "GPS-Track mit 2.284 Messpunkten sowie zwischengespeicherte Wetter- und Ortsdaten."),
            ("Berechnung: ", "Fahrdynamik, Umwelteinflüsse, Motorbedarf und zwei Akkuvarianten werden abschnittsweise simuliert."),
            ("Ausgabe: ", "Kennzahlen im Terminal, 21 auswählbare Diagramme, eine HTML-Karte und optional ein PDF-Ergebnisbericht."),
            ("Vergleich: ", "Neben einer Einzelsimulation steht eine Parameterstudie mit 20 vordefinierten Szenarien zur Verfügung."),
        ],
    )

    document.add_heading("1.2 Umfang und Modellcharakter", level=2)
    add_body(
        document,
        "Das Projekt bildet eine reale Fahrt mit vereinfachten physikalischen Modellen nach. Es handelt sich nicht um eine vollständige Auslegung eines serienreifen E-Bike-Antriebs. Ziel ist vielmehr, die Verarbeitung von Messdaten, objektorientierte Programmierung, numerische Berechnungen, Visualisierung und automatisierte Tests in einem zusammenhängenden Projekt zu verbinden."
    )
    add_callout(
        document,
        "Referenzlauf",
        "Mit den Standardwerten verarbeitet das Programm die enthaltene Route über 94,33 km in 272,9 Minuten. Der berechnete Endladezustand beträgt für beide Akkuvarianten rund 28,9 Prozent. Diese Werte dienen in der Dokumentation als überprüfbares Beispiel und nicht als allgemeingültige Reichweitenangabe.",
    )

    document.add_heading("2 Projekt in Betrieb nehmen", level=1)
    document.add_heading("2.1 Voraussetzungen", level=2)
    add_body(
        document,
        "Für die Ausführung wird eine aktuelle Python-Version benötigt. Das vorhandene virtuelle Environment wurde mit Python 3.14.3 erstellt; der verwendete Quellcode setzt mindestens eine moderne Python-3-Version voraus. Die folgenden Schritte werden aus dem Ordner ebike-simulation ausgeführt."
    )
    add_rich_list(
        document,
        [
            ("Betriebssystem: ", "Windows, Linux oder macOS mit Terminal beziehungsweise PowerShell."),
            ("Python: ", "Python 3.10 oder neuer wird empfohlen; pip muss verfügbar sein."),
            ("Eingabedatei: ", "data/final_project_input_data.csv mit den Spalten lat, lon, ele, time und temperature."),
            ("Internet: ", "Nur erforderlich, wenn für Koordinaten oder Zeitpunkte noch keine Wetter- oder Ortsdaten im Cache vorhanden sind."),
        ],
    )

    document.add_heading("2.2 Installation unter Windows", level=2)
    add_list(
        document,
        [
            "PowerShell öffnen und in den Projektordner wechseln.",
            "Ein virtuelles Python-Environment erstellen.",
            "Das Environment aktivieren und die Abhängigkeiten installieren.",
            "Das Programm mit main.py starten.",
        ],
        kind="decimal",
    )
    add_code_block(
        document,
        "cd ebike-simulation\npython -m venv .venv\n.\\.venv\\Scripts\\Activate.ps1\npython -m pip install --upgrade pip\npython -m pip install -r requirements.txt\npython -m pip install requests\npython main.py",
    )
    add_callout(
        document,
        "Abhängigkeit",
        "Das Modul requests wird von get_weather_data.py und reverse_geocoding.py importiert, ist in der aktuellen requirements.txt jedoch noch nicht aufgeführt. Daher wird es im gezeigten Installationsablauf zusätzlich installiert.",
        color=ORANGE,
    )

    document.add_heading("2.3 Installation unter Linux oder macOS", level=2)
    add_code_block(
        document,
        "cd ebike-simulation\npython3 -m venv .venv\nsource .venv/bin/activate\npython -m pip install --upgrade pip\npython -m pip install -r requirements.txt\npython -m pip install requests\npython main.py",
    )

    document.add_heading("2.4 Verwendete Bibliotheken", level=2)
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Bibliothek"
    table.rows[0].cells[1].text = "Aufgabe im Projekt"
    for library, purpose in (
        ("NumPy", "Vektorisierte numerische Berechnungen und Datenreihen"),
        ("pandas", "Einlesen und Verarbeiten der CSV-, Zeit- und Wetterdaten"),
        ("Matplotlib", "Erstellung der auswählbaren Ergebnisdiagramme"),
        ("ReportLab", "Erzeugung des optionalen PDF-Ergebnisberichts"),
        ("Folium", "Erstellung der interaktiven HTML-Karte der Route"),
        ("requests", "HTTP-Zugriff auf Open-Meteo und Geoapify bei Cache-Fehlstellen"),
    ):
        cells = table.add_row().cells
        cells[0].text = library
        cells[1].text = purpose
    style_table(table, [2100, 6972])

    document.add_heading("2.5 Eingaben, Ausgaben und Tests", level=2)
    add_body(
        document,
        "Die GPS-Datei liegt fest unter data/final_project_input_data.csv. Beim Start wird sie relativ zu main.py aufgelöst. Während der Datenvorbereitung erzeugt GPSMap immer outputs/karte.html. Wenn im Menü ein PDF-Bericht gewählt wird, entsteht zusätzlich outputs/ebike_simulation_report.pdf. Diagramme werden als Matplotlib-Fenster angezeigt. Wetter- und Ortsdaten werden in data/weather_cache.json beziehungsweise data/geocode_cache.json wiederverwendet."
    )
    add_body(document, "Die automatisierten Tests werden mit folgendem Befehl gestartet:")
    add_code_block(document, "python -m unittest discover -s tests -v")
    add_body(document, "Beim dokumentierten Prüflauf am 19.07.2026 wurden alle 20 Tests erfolgreich ausgeführt.")

    document.add_heading("3 Bedienung des Terminalmenüs", level=1)
    document.add_heading("3.1 Auswahl der Betriebsart", level=2)
    add_body(
        document,
        "Nach dem Start fragt das Programm zuerst, ob eine Parameterstudie oder eine Simulation mit konkreten Werten durchgeführt werden soll. Die Eingabe 1 startet die Parameterstudie. Dabei werden 20 physikalisch sinnvolle Szenarien berechnet, unter anderem Pendler-, Lastenrad-, Sport-, MTB-, Massen-, Luftwiderstands-, Reifen- und Antriebsvarianten sowie Grenzfälle. Die Ausgabe zeigt den LiPo-Endladezustand und seine Abweichung vom Basisfall in Prozentpunkten und Prozent."
    )
    add_body(
        document,
        "Die Eingabe 2 startet eine einzelne, interaktiv konfigurierte Simulation. Hier können Parameter, Terminalkennzahlen, Diagramme und der PDF-Bericht ausgewählt werden. Diese Betriebsart eignet sich für einen gezielten Versuch oder für die Demonstration des Programms."
    )

    document.add_heading("3.2 Simulationsparameter", level=2)
    add_body(
        document,
        "Im Einzelsimulationsmodus kann die Eingabetaste alle Standardwerte übernehmen. Wird n gewählt, fragt das Programm die folgenden sechs Werte nacheinander ab. Eine leere Eingabe übernimmt jeweils den angezeigten Standardwert; Dezimalzahlen können mit Punkt oder Komma eingegeben werden. Alle Werte müssen positiv sein."
    )
    table = document.add_table(rows=1, cols=4)
    for index, value in enumerate(("Parameter", "Bedeutung", "Standard", "Einheit")):
        table.rows[0].cells[index].text = value
    rows = (
        ("rider_mass_kg", "Fahrergewicht", "70,0", "kg"),
        ("bike_mass_kg", "Fahrradgewicht", "10,0", "kg"),
        ("drag_area_m2", "Effektive Stirnfläche cw·A", "0,5625", "m²"),
        ("wheel_diameter_inch", "Raddurchmesser", "27,0", "Zoll"),
        ("motor_constant_nm_per_a", "Motorkonstante", "1,5", "Nm/A"),
        ("rolling_resistance_coefficient", "Rollwiderstandsbeiwert", "0,0077", "-"),
    )
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            if index >= 2:
                cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_table(table, [2550, 3572, 1550, 1400])

    document.add_heading("3.3 Auswahlregeln in den Menüs", level=2)
    add_rich_list(
        document,
        [
            ("Einzelne Einträge: ", "Nummer eingeben, zum Beispiel 3."),
            ("Mehrere Einträge: ", "Nummern durch Komma oder Leerzeichen trennen, zum Beispiel 1,3,5 oder 1 3 5."),
            ("Alle Einträge: ", "a oder alle eingeben."),
            ("Keine Einträge: ", "0 oder keine eingeben."),
            ("Programm abbrechen: ", "q, quit oder abbrechen eingeben. Dies ist aus jedem Auswahlmenü möglich."),
        ],
    )
    add_body(
        document,
        "Ungültige Texte, Zahlen außerhalb des angebotenen Bereichs und leere Auswahlen werden abgefangen. Doppelte Nummern erscheinen nur einmal in der internen Auswahl. Ja-Nein-Fragen akzeptieren j/ja, n/nein und q zum Abbrechen."
    )

    document.add_heading("3.4 Terminalkennzahlen", level=2)
    add_body(document, "Für die Konsolenausgabe stehen fünf Ergebnisgruppen zur Verfügung:")
    add_rich_list(
        document,
        [
            ("Route und Fahrdaten: ", "Gesamtstrecke, Fahrzeit, mittlere und maximale Geschwindigkeit, Aufstieg und Abstieg."),
            ("Umgebungsdaten: ", "Temperaturbereich sowie mittlere, minimale und maximale Luftdichte."),
            ("Motor und Antrieb: ", "Rollwiderstand, maximale Leistung und Stromstärke sowie mechanische Energie."),
            ("Akkudaten: ", "Endladezustand, Minimalspannung, Anfangs-, Maximal- und Endtemperaturen sowie maximale Akkuleistung für LiPo und NMC."),
            ("Rekuperation und Bremsen: ", "Bremsenergie, zurückgewonnene Energie, Widerstandsenergie, mechanische Bremsenergie, Leistungen und Temperaturen."),
        ],
    )

    document.add_heading("3.5 Diagramme und PDF-Bericht", level=2)
    add_body(document, "Die 21 Diagrammoptionen sind thematisch wie folgt gegliedert:")
    add_rich_list(
        document,
        [
            ("Fahrt und Umgebung: ", "Geschwindigkeit, Beschleunigung, Steigung, Luftdichte, zurückgelegte Strecke und Windkraft."),
            ("Motor: ", "Antriebskraft, Motorleistung, Drehmoment und Motorstrom."),
            ("Akku: ", "Akkustrom, SoC, Spannung, Temperatur, Leistung und Innenwiderstand."),
            ("Bremsen: ", "Bremsleistungsbedarf, Rekuperationsleistung, Bremswiderstandsleistung, Bremswiderstandstemperatur und mechanische Bremsleistung."),
        ],
    )
    add_body(
        document,
        "Nach der Diagrammauswahl fragt das Programm, ob aus derselben Auswahl ein PDF-Bericht erstellt werden soll. Der PDF-Bericht enthält die gewählten Kennzahlengruppen und für jedes gewählte Diagramm eine eigene Seite. Er wird vor dem Öffnen der Diagrammfenster geschrieben, damit blockierende Fenster die Dateierstellung nicht verzögern."
    )

    document.add_heading("3.6 Beispiel für eine sinnvolle Menüfolge", level=2)
    add_code_block(
        document,
        "2        # konkrete Werte simulieren\n[Enter]  # alle Standardparameter übernehmen\n1,4,5    # Route, Akkus und Rekuperation im Terminal\n8,12,18  # Motorleistung, SoC und Rekuperationsleistung\nj        # PDF-Bericht erzeugen",
    )
    add_body(
        document,
        "Nach Abschluss liegen die Karte und der PDF-Bericht im Ordner outputs. Die drei ausgewählten Diagramme werden anschließend nacheinander angezeigt. Bei q wird kontrolliert beendet; bei Strg+C endet das Programm mit dem Rückgabecode 130. Datei-, Daten- und andere Laufzeitfehler werden über das Logging ausgegeben."
    )

    document.add_page_break()
    document.add_heading("4 Softwarearchitektur und UML", level=1)
    document.add_heading("4.1 Architekturprinzip", level=2)
    add_body(
        document,
        "Die Architektur folgt einer klaren Aufteilung in Programmsteuerung, Datenaufbereitung, physikalische Modelle und Ausgabe. main.py ist die Benutzerschnittstelle. Die Klasse BikeSimulator übernimmt die Rolle eines Orchestrators: Sie kennt den vollständigen Ablauf, überlässt die fachlichen Teilaufgaben aber spezialisierten Klassen und Funktionen. Dadurch können etwa RouteCalculator, Motor oder BatteryPack unabhängig getestet werden."
    )
    add_body(
        document,
        "Der Datenaustausch erfolgt überwiegend über NumPy-Arrays und verschachtelte Dictionaries. Zeitabhängige Werte besitzen pro Streckenabschnitt dieselbe Länge. Das abschließende Ergebnis-Dictionary trennt die Bereiche time, route, environment, motor, battery und braking und enthält zusätzlich die zusammengefassten metrics. Diese gemeinsame Struktur wird von Konsole, Plotter und PDF-Bericht verwendet."
    )
    add_figure(
        document,
        UML_FILE,
        "Abbildung 1: Vereinfachtes UML-Klassendiagramm der E-Bike-Simulation",
        "UML-Klassendiagramm mit main.py und BikeSimulator als Steuerung, Datenklassen links, Akku-Vererbung in der Mitte sowie Motor- und Bremsklassen rechts.",
        6.15,
    )

    document.add_heading("4.2 Beziehungen im Diagramm", level=2)
    add_rich_list(
        document,
        [
            ("Aufrufbeziehung: ", "main.py erzeugt BikeSimulator und ruft run() auf. Die Ergebnisstruktur wird danach an die Ausgabemodule weitergegeben."),
            ("Abhängigkeiten: ", "BikeSimulator verwendet die Daten-, Motor-, Akku- und Bremsklassen, ohne deren interne Berechnungen zu duplizieren."),
            ("Vererbung: ", "BatteryBase definiert die abstrakte Schnittstelle. BatteryPack implementiert das gemeinsame elektrische und thermische Verhalten. LiPoBatteryPack und NMCBatteryPack überschreiben die Spannungskennlinie."),
            ("Komposition: ", "Für jeden Simulationslauf erstellt BikeSimulator eigenständige Akku- und Bremswiderstandsobjekte. LiPo und NMC werden dadurch mit identischen Eingangsdaten, aber getrennten Zuständen simuliert."),
        ],
    )

    document.add_heading("5 Module des Projekts", level=1)
    document.add_heading("5.1 Programmsteuerung", level=2)
    document.add_heading("main.py", level=3)
    add_body(
        document,
        "main.py ist der Einstiegspunkt des Programms. Es definiert Dateipfade, Standardparameter, 20 Parametersätze und sämtliche Terminaldialoge. Die Hilfsfunktionen select_items(), ask_yes_no(), ask_float() und ask_parameters() prüfen Benutzereingaben. run_study() führt die Szenarien aus, add_basis_vergleich() ergänzt Abweichungen zum Basisfall und main() steuert Fehlerbehandlung, Logging und Ausgabe."
    )

    document.add_heading("src/bikesimulator.py", level=3)
    add_body(
        document,
        "BikeSimulator ist die zentrale Anwendungsklasse. Der Konstruktor speichert und validiert alle Parameter. Die interne Verarbeitung ist in fünf Phasen gegliedert: Routendaten vorbereiten, Motordaten berechnen, Simulationskomponenten erzeugen, beide Akkuvarianten abschnittsweise simulieren und die Kennzahlen beziehungsweise die Ergebnisstruktur aufbauen. Die öffentliche Methode run() verbindet diese Phasen."
    )

    document.add_heading("5.2 Datenaufbereitung und Umgebungsdaten", level=2)
    document.add_heading("src/gps_reader.py", level=3)
    add_body(
        document,
        "GPSReader liest die semikolongetrennte CSV-Datei mit pandas ein, kontrolliert die erforderlichen Spalten und fehlende Werte und berechnet anschließend die dreidimensionalen Abstände. Für den horizontalen Anteil verwendet die Klasse die Haversine-Formel, für die 3D-Distanz zusätzlich den Höhenunterschied. Aufstieg, Abstieg und Gesamtstrecke werden als Objektzustand gespeichert."
    )

    document.add_heading("src/route_calculator.py", level=3)
    add_body(
        document,
        "RouteCalculator berechnet Geschwindigkeit, Beschleunigung und Steigung. Die Geschwindigkeit ergibt sich aus Distanz und Zeitintervall. Die Beschleunigung wird aus der Änderung aufeinanderfolgender Geschwindigkeiten bestimmt und anschließend mit einem Nachbarschaftsverfahren geglättet. Die Steigung ist das Verhältnis aus Höhenänderung und räumlicher Abschnittslänge. Jede Methode validiert Längen, Zahlenwerte und physikalisch unzulässige Eingaben."
    )

    document.add_heading("src/air_density.py", level=3)
    add_body(
        document,
        "calculate_air_density() ist eine eigenständige Funktion. Sie bestimmt zunächst den höhenabhängigen Luftdruck nach einem vereinfachten Standardatmosphärenmodell und berechnet danach mit der idealen Gasgleichung die Luftdichte für jedes Temperatur-Höhen-Paar."
    )

    document.add_heading("src/get_driving_direction.py", level=3)
    add_body(
        document,
        "MovementDirection bestimmt aus je zwei aufeinanderfolgenden Koordinaten den Kurswinkel zwischen 0 und 360 Grad. Dieser Bearing-Wert wird benötigt, um die Windrichtung in einen längs zur Fahrtrichtung wirkenden Anteil umzuwandeln."
    )

    document.add_heading("src/get_weather_data.py", level=3)
    add_body(
        document,
        "TripWeather ordnet jedem GPS-Punkt Wetterdaten zu. Koordinaten werden auf zwei Dezimalstellen und Zeitpunkte auf 15-Minuten-Intervalle gerundet. Bereits vorhandene Kombinationen werden aus weather_cache.json gelesen; nur fehlende Zellen fragt die Klasse gesammelt bei Open-Meteo ab. Verwendet werden Temperatur, Luftdruck, Windgeschwindigkeit, Windrichtung, Böen und Niederschlag. In der aktuellen Motorsimulation fließen vor allem Windgeschwindigkeit und Windrichtung ein."
    )

    document.add_heading("src/reverse_geocoding.py", level=3)
    add_body(
        document,
        "Reverse_Geocoder rundet Koordinaten auf drei Dezimalstellen und ergänzt Orts- und Straßennamen über den Geoapify-Batchdienst. geocode_cache.json verhindert wiederholte Anfragen. Die Ortsnamen werden später für die farbliche Gliederung einiger Diagramme verwendet. Für eine Veröffentlichung sollte der aktuell als Standardwert hinterlegte API-Schlüssel durch eine Umgebungsvariable ersetzt werden."
    )

    document.add_heading("src/data_cleaner.py", level=3)
    add_body(
        document,
        "Cleaner vereinheitlicht bekannte Ortsteilbezeichnungen und entfernt sehr kurze, wahrscheinlich durch Geocoding-Unschärfen entstandene Ortswechsel. Kleine Abschnitte werden so lange dem größeren Nachbarabschnitt zugeordnet, bis die Mindestlänge erreicht ist."
    )

    document.add_heading("src/gps_plot_route_on_map.py", level=3)
    add_body(
        document,
        "GPSMap erzeugt mit Folium eine interaktive Karte. Die Route wird als Linie dargestellt, Start und Ziel erhalten Marker; liegen beide weniger als zehn Meter auseinander, wird ein gemeinsamer Start-/Zielmarker verwendet. save() schreibt die Datei standardmäßig nach outputs/karte.html."
    )

    document.add_heading("5.3 Physikalische Modelle", level=2)
    document.add_heading("src/motor.py", level=3)
    add_body(
        document,
        "Motor bündelt das Fahrwiderstands- und Antriebsmodell. Berechnet werden Beschleunigungs-, Steigungs-, Roll- und Luftkraft einschließlich scheinbarem Wind. Aus der Gesamtkraft folgen Leistung und Drehmoment. Positive Werte werden als Antrieb, negative Werte als Bremsbedarf interpretiert. Über die Motorkonstante wird aus dem positiven Drehmoment ein Antriebsstrom bestimmt."
    )

    document.add_heading("src/battery_base.py und src/battery_pack.py", level=3)
    add_body(
        document,
        "BatteryBase ist eine abstrakte Basisklasse für apply_current() und voltage(). BatteryPack implementiert die gemeinsame Zustandslogik. Die Klasse begrenzt den SoC auf 0 bis 100 Prozent, berücksichtigt maximale Lade-, Spannungs- und SoC-Grenzen und berechnet den temperaturabhängigen Innenwiderstand. Ein thermisches Ersatzmodell verbindet I²R-Verluste, Wärmeabgabe und thermische Kapazität."
    )

    document.add_heading("src/lipo_battery.py und src/nmc_battery.py", level=3)
    add_body(
        document,
        "LiPoBatteryPack und NMCBatteryPack erben von BatteryPack. Beide übernehmen Kapazitäts-, Strom-, Widerstands- und Temperaturmodell, verwenden aber unterschiedliche nichtlineare Leerlaufspannungskennlinien. Zusätzlich startet das LiPo-Modell mit 80 mΩ Referenzinnenwiderstand, das NMC-Modell mit 70 mΩ."
    )

    document.add_heading("src/regenerative_braking.py", level=3)
    add_body(
        document,
        "RegenerativeBrakingController verteilt den Bremsleistungsbedarf. Im Modell können 75 Prozent der mechanischen Bremsleistung, maximal jedoch 500 W, elektrisch bereitgestellt werden. Zuerst wird der zulässige Ladestrom des Akkus ausgeschöpft. Danach folgt der Bremswiderstand. Den Rest übernimmt die mechanische Bremse. Die Methode liefert zusätzlich die Umwandlungsverluste und garantiert eine nachvollziehbare Leistungsbilanz."
    )

    document.add_heading("src/brake_resistor.py", level=3)
    add_body(
        document,
        "BrakeResistor begrenzt die aufnehmbare Leistung sowohl über P = U²/R als auch über seine Nennleistung. Ein thermisches Modell berechnet Erwärmung und Abkühlung gegenüber der Umgebung. Die insgesamt in Wärme umgesetzte Energie wird in Joule gesammelt und über eine Property in Wattstunden bereitgestellt."
    )

    document.add_heading("5.4 Ausgabe und Berichte", level=2)
    document.add_heading("src/plotter.py", level=3)
    add_body(
        document,
        "plotter.py definiert die 21 auswählbaren Diagramme und deren Datenreihen, Achsen und Beschriftungen. create_result_figure() erzeugt eine einzelne Matplotlib-Figur, show_result_figures() zeigt die gewählte Menge an. Ortsabschnitte können farblich unterschieden werden. Für LiPo und NMC werden gemeinsame Achsen verwendet, sodass Unterschiede direkt sichtbar sind."
    )

    document.add_heading("src/reporting/console.py", level=3)
    add_body(
        document,
        "console.py formatiert die Kennzahlengruppen als Textzeilen. Dieselben Formatierungsfunktionen werden sowohl für die Terminalausgabe als auch für den PDF-Bericht verwendet. Dadurch bleiben Bezeichnungen und Einheiten konsistent. print_vergleich() übernimmt die kompakte Ausgabe der Parameterstudie."
    )

    document.add_heading("src/reporting/pdf_report.py", level=3)
    add_body(
        document,
        "pdf_report.py erstellt mit ReportLab einen A4-Bericht. Gewählte Kennzahlen werden als Textabschnitte formatiert, Diagramme werden im Arbeitsspeicher als PNG gerendert und jeweils auf einer eigenen PDF-Seite eingefügt. Seitennummern und Metadaten werden automatisch ergänzt."
    )

    document.add_heading("Paketdateien __init__.py", level=3)
    add_body(
        document,
        "Die Dateien src/__init__.py und src/reporting/__init__.py kennzeichnen die Verzeichnisse als Python-Pakete. Sie enthalten aktuell keine weitere Programmlogik."
    )

    document.add_page_break()
    document.add_heading("6 Erklärung der Berechnungsmodelle", level=1)
    document.add_heading("6.1 Zentrale Verarbeitung in BikeSimulator", level=2)
    add_body(
        document,
        "Die Methode run() enthält bewusst wenig Detailmathematik. Stattdessen ruft sie die fachlichen Teilschritte in einer klaren Reihenfolge auf. In vereinfachter Form entspricht der Ablauf folgendem Pseudocode:"
    )
    add_code_block(
        document,
        "route_data = prepare_route_data()\nmotor_data = calculate_motor_data(route_data)\ncomponents = create_simulation_components()\nlipo_data = simulate_battery_variant(LiPo, route_data, motor_data)\nnmc_data = simulate_battery_variant(NMC, route_data, motor_data)\nmetrics = calculate_metrics(route_data, motor_data, lipo_data, nmc_data)\nreturn build_results(metrics, route_data, motor_data, lipo_data, nmc_data)",
    )
    add_body(
        document,
        "Diese Zerlegung erleichtert das Lesen und Testen. Die ausführliche numerische Verarbeitung bleibt in den jeweiligen Modulen, während BikeSimulator nur die Reihenfolge und den Datenaustausch festlegt."
    )

    document.add_heading("6.2 Route und Bewegung", level=2)
    add_rich_list(
        document,
        [
            ("Distanz: ", "Die horizontale Distanz wird mit der Haversine-Formel bestimmt. Zusammen mit dem Höhenunterschied ergibt sich d3D = √(dhorizontal² + Δh²)."),
            ("Geschwindigkeit: ", "Für jeden Abschnitt gilt v = s / Δt. Null oder negative Zeitintervalle werden als ungültig abgelehnt."),
            ("Beschleunigung: ", "a = (vi - vi-1) / Δt. Auffällige Werte werden mit dem Mittelwert benachbarter Werte geglättet."),
            ("Steigung: ", "slope = Δh / s. Im Motormodell entspricht dieser Wert dem Sinus des Steigungswinkels."),
            ("Luftdichte: ", "Aus Höhe und Temperatur werden zunächst Luftdruck und danach ρ = p / (R · T) berechnet."),
        ],
    )

    document.add_heading("6.3 Motor und Fahrwiderstände", level=2)
    add_body(
        document,
        "Die Gesamtkraft ist die Summe aus Beschleunigungs-, Steigungs-, Roll- und Luftkraft. Vereinfacht gilt: Fges = m·a + m·g·slope + crr·m·g·cos(α) + FLuft. Für die Luftkraft werden Fahrtrichtung, Windrichtung und Windgeschwindigkeit zu einer scheinbaren Windgeschwindigkeit kombiniert. Dadurch kann Gegenwind den Bedarf erhöhen und Rückenwind ihn senken."
    )
    add_body(
        document,
        "Aus der Kraft folgen die wichtigsten Antriebsgrößen: P = Fges·v, M = Fges·rRad und I = M / kM. Ist die vorzeichenbehaftete Leistung positiv, wird sie als Motorleistung verwendet. Ein negativer Wert bedeutet, dass zum Einhalten des gemessenen Fahrprofils gebremst werden muss; sein Betrag wird an die Rekuperationssimulation weitergegeben."
    )

    document.add_heading("6.4 Akku, Spannung und Temperatur", level=2)
    add_body(
        document,
        "Der Ladezustand wird über die Strombilanz aktualisiert: SoCneu = clamp(SoCalt - I·Δt / Cnom, 0, 1). Ein positiver Strom entlädt den Akku, ein negativer Strom lädt ihn. Die Leerlaufspannung liegt zwischen Vmin und Vmax. Für LiPo wird der SoC mit dem Exponenten 0,3, für NMC mit 0,834 in die Kennlinie eingesetzt. Unter Last gilt UKlemme = UOCV - I·R(T)."
    )
    add_body(
        document,
        "Der Innenwiderstand steigt im vereinfachten Modell bei niedriger Temperatur und wird zwischen dem 0,5- und 3-fachen Referenzwert begrenzt. Die Wärmeleistung im Akku beträgt I²·R. Gleichzeitig wird Wärme über den thermischen Widerstand an die Umgebung abgegeben. Aus Nettoleistung, Dauer und thermischer Kapazität folgt die Temperaturänderung."
    )

    document.add_heading("6.5 Rekuperation und Bremsleistung", level=2)
    add_body(
        document,
        "In einem Bremsabschnitt begrenzen der Wirkungsgrad von 75 Prozent und die Generatorgrenze von 500 W das elektrische Potenzial: Pelektrisch = min(Pbrems·0,75; 500 W). Der Controller berechnet daraus einen möglichen Ladestrom. BatteryPack begrenzt diesen Wert anhand von maximalem Ladestrom, Spannungsreserve, aktuellem SoC und Dauer des Zeitschritts."
    )
    add_body(
        document,
        "Nicht vom Akku aufgenommene elektrische Leistung wird dem Bremswiderstand angeboten. Dessen Grenze folgt aus Spannung, Widerstand und Nennleistung. Alles, was danach noch nicht verarbeitet werden kann, bleibt mechanische Bremsleistung. Dieser priorisierte Ablauf erklärt, warum ein voller Akku stärker auf Bremswiderstand und mechanische Bremse angewiesen ist."
    )

    document.add_heading("6.6 Ergebnisstruktur und Parameterstudie", level=2)
    add_body(
        document,
        "Nach beiden Akkuvarianten werden Energien durch Summieren von Leistung mal Zeit berechnet und von Wattsekunden in Wattstunden umgerechnet. Die Ergebnisstruktur enthält sowohl vollständige Zeitreihen als auch kompakte Kennzahlen. Bei der Parameterstudie wird derselbe Simulationskern 20-mal mit abweichenden Massen-, Luftwiderstands-, Reifen- oder Antriebsparametern ausgeführt. Anschließend werden die Endladezustände mit dem Basisfall verglichen."
    )

    document.add_page_break()
    document.add_heading("7 Ablauf der Simulation", level=1)
    document.add_heading("7.1 Aktivitätsdiagramm", level=2)
    add_figure(
        document,
        ACTIVITY_FILE,
        "Abbildung 2: Aktivitätsdiagramm der Einzelsimulation",
        "Aktivitätsdiagramm vom Erzeugen des BikeSimulator über Routendaten, Motormodell, Schleifen für Akkuvarianten und Streckenabschnitte bis zur Ausgabe.",
        5.95,
    )

    document.add_heading("7.2 Erläuterung des Ablaufs", level=2)
    add_list(
        document,
        [
            "Nach der Menüauswahl wird BikeSimulator mit Dateipfad, Akkuvorgaben und den gewählten Fahrradparametern erzeugt. Ungültige Grundparameter führen sofort zu einem ValueError.",
            "GPS- und Umgebungsdaten werden vorbereitet. Dabei entstehen auch die HTML-Karte sowie gegebenenfalls neue Cache-Einträge für Wetter und Geocoding.",
            "Das Motormodell berechnet für jeden Streckenabschnitt Antriebs- oder Bremsbedarf.",
            "LiPo, NMC, Rekuperationscontroller und zwei unabhängige Bremswiderstände werden mit der ersten Umgebungstemperatur initialisiert.",
            "Für LiPo und NMC wird derselbe Streckenverlauf durchlaufen. Jeder Abschnitt entscheidet anhand der Bremsleistung zwischen Antriebs- und Rekuperationsfall.",
            "Nach der Stromwahl werden SoC und Akkutemperatur aktualisiert. Der Bremswiderstand erwärmt sich bei Leistungsaufnahme und kühlt sonst in Richtung Umgebungstemperatur ab.",
            "Alle Zeitreihen werden gespeichert. Nach beiden Akkuvarianten werden Kennzahlen, Energiebilanzen und das gemeinsame Ergebnis-Dictionary aufgebaut.",
            "main.py gibt nur die gewählten Kennzahlen aus, erzeugt auf Wunsch den PDF-Bericht und zeigt zuletzt die gewählten Diagramme.",
        ],
        kind="decimal",
    )

    document.add_heading("8 Verifikation, Grenzen und Verbesserungen", level=1)
    document.add_heading("8.1 Automatisierte Tests", level=2)
    add_body(
        document,
        "Die Tests verwenden das Python-Modul unittest. Sie konzentrieren sich auf berechenbare Kernfunktionen und Grenzfälle. Der dokumentierte Lauf umfasste 20 Tests; alle wurden bestanden."
    )
    table = document.add_table(rows=1, cols=3)
    for index, value in enumerate(("Testdatei", "Schwerpunkt", "Ergebnis")):
        table.rows[0].cells[index].text = value
    for row in (
        ("test_route_calculator.py", "Geschwindigkeit, Beschleunigung, Steigung und ungültige Eingaben", "6 bestanden"),
        ("test_motor.py", "Antrieb, Bremsfall und Eingabevalidierung", "3 bestanden"),
        ("test_battery_pack.py", "Entladen, Laden, Ladestromgrenzen und Fehlerfälle", "4 bestanden"),
        ("test_regenerative_braking.py", "Leistungsverteilung, voller Akku, Nullfälle und Bilanz", "4 bestanden"),
        ("test_air_density.py", "Standardatmosphäre, Trends und ungültige Werte", "3 bestanden"),
    ):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_table(table, [2600, 4772, 1700])

    document.add_heading("8.2 Ergebnisse des verifizierten Referenzlaufs", level=2)
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Kennzahl"
    table.rows[0].cells[1].text = "Standardlauf"
    for metric, value in (
        ("Strecke / Fahrzeit", "94,33 km / 272,9 min"),
        ("Durchschnitt / Maximum", "20,7 km/h / 49,2 km/h"),
        ("Aufstieg / Abstieg", "1.096 m / 1.097 m"),
        ("Maximale Motorleistung", "916 W"),
        ("Mechanische Antriebsenergie", "724,4 Wh"),
        ("Mechanische Bremsenergie", "34,34 Wh"),
        ("Zurückgewonnene Energie", "LiPo 23,01 Wh / NMC 23,21 Wh"),
        ("Endladezustand", "LiPo 28,9 % / NMC 28,9 %"),
        ("Maximale Akkutemperatur", "LiPo 34,94 °C / NMC 34,15 °C"),
    ):
        cells = table.add_row().cells
        cells[0].text = metric
        cells[1].text = value
    style_table(table, [3950, 5122])
    add_callout(
        document,
        "Interpretation",
        "Die NMC-Variante besitzt im Referenzlauf die niedrigere Minimalspannung und etwas geringere Maximaltemperatur. Die Endladezustände liegen trotz unterschiedlicher Spannungsverläufe nahezu gleichauf. Für belastbare Aussagen zu realen Akkus wären Messdaten und eine Kalibrierung der Modellparameter erforderlich.",
    )

    document.add_heading("8.3 Grenzen des Modells", level=2)
    add_rich_list(
        document,
        [
            ("Messdaten: ", "GPS-Höhe, Zeitintervalle und Temperatur enthalten Messrauschen. Die Beschleunigungsglättung reduziert, beseitigt aber nicht alle Ausreißer."),
            ("Antrieb: ", "Motor- und Leistungselektronikverluste werden nur teilweise abgebildet; der Motorstrom folgt direkt aus Drehmoment und Motorkonstante."),
            ("Akku: ", "OCV-Kennlinien, Innenwiderstand und thermische Parameter sind vereinfachte Annahmen und nicht anhand eines konkreten Packs kalibriert."),
            ("Rekuperation: ", "Wirkungsgrad, Generatorgrenze und Bremswiderstand sind feste Modellparameter."),
            ("Externe Dienste: ", "Fehlende Cache-Daten benötigen Internetzugriff und funktionierende APIs. Änderungen an Diensten oder Limits können den Abruf beeinflussen."),
        ],
    )

    document.add_heading("8.4 Sinnvolle nächste Verbesserungen", level=2)
    add_list(
        document,
        [
            "requests in requirements.txt ergänzen und eine getestete Python-Mindestversion im README angeben.",
            "Den Geoapify-API-Schlüssel ausschließlich über GEOAPIFY_API_KEY einlesen und keinen Schlüssel im Quellcode speichern.",
            "GPS-Dateipfad, Akkukapazität, Start-SoC, Filterfenster und Rekuperationsparameter über das Menü oder eine Konfigurationsdatei einstellbar machen.",
            "Integrationstests für einen vollständigen Lauf mit kleinen Testdaten und vorbereiteten Cache-Dateien ergänzen.",
            "Messdaten realer Akkus verwenden, um Spannungskennlinien, Innenwiderstand und thermisches Verhalten zu kalibrieren.",
            "Die Parameterstudie zusätzlich als CSV oder PDF exportieren und LiPo sowie NMC in derselben Vergleichstabelle ausgeben.",
        ],
        kind="decimal",
    )

    document.add_page_break()
    document.add_heading("9 Verwendung von künstlicher Intelligenz", level=1)
    document.add_heading("9.1 Unterstützende Rolle", level=2)
    add_body(
        document,
        "Generative künstliche Intelligenz wurde bei der Erstellung dieser Projektdokumentation als unterstützendes Werkzeug eingesetzt. Sie half dabei, die vorhandene Codebasis systematisch zu sichten, eine nachvollziehbare Gliederung zu entwerfen, die Beziehungen für das UML-Diagramm und den Ablauf für das Aktivitätsdiagramm aus dem Quellcode abzuleiten sowie Formulierungen sprachlich zu überarbeiten."
    )
    add_body(
        document,
        "Die KI ersetzte nicht die fachliche Prüfung. Alle technischen Aussagen wurden mit den vorhandenen Python-Dateien abgeglichen. Zusätzlich wurden die 20 Unit-Tests und ein vollständiger Programmlauf mit den Standardwerten ausgeführt. Zahlen im Bericht stammen aus diesem Lauf oder direkt aus den im Code festgelegten Parametern."
    )

    document.add_heading("9.2 Verantwortlicher Umgang", level=2)
    add_rich_list(
        document,
        [
            ("Prüfung: ", "Vorschläge der KI wurden nicht ungeprüft übernommen, sondern mit Code, Menüausgabe und Testergebnissen verglichen."),
            ("Transparenz: ", "Der Einsatz wird im Bericht offengelegt. KI-Unterstützung und eigene technische Verantwortung werden getrennt dargestellt."),
            ("Datenschutz: ", "Für die Dokumentation wurden ausschließlich Projektdateien und technische Daten verwendet; personenbezogene Daten sind nicht Bestandteil der beschriebenen Verarbeitung."),
            ("Urheberschaft: ", "Entscheidungen über Inhalt, Auswahl, Bewertung und endgültige Abgabe bleiben bei der Autorin beziehungsweise beim Autor des Projekts."),
        ],
    )
    add_callout(
        document,
        "Vor der Abgabe prüfen",
        "Wenn KI bereits während der Programmierung genutzt wurde, sollten hier zusätzlich das verwendete Werkzeug und konkrete Aufgaben ergänzt werden, zum Beispiel Fehlersuche, Erklärung einer Bibliothek oder Entwurf einzelner Tests. Nur tatsächlich erfolgte Nutzung angeben.",
        color=ORANGE,
    )

    document.add_heading("10 Fazit", level=1)
    add_body(
        document,
        "Das Abschlussprojekt verbindet Datenverarbeitung, physikalische Modellierung und objektorientierte Programmierung zu einer durchgängigen E-Bike-Simulation. Die Aufteilung in spezialisierte Module macht den Programmablauf verständlich und ermöglicht gezielte Tests der zentralen Berechnungen. Besonders die parallele Simulation von LiPo und NMC sowie die Aufteilung der Bremsenergie auf Akku, Widerstand und mechanische Bremse gehen über eine reine GPS-Auswertung hinaus."
    )
    add_body(
        document,
        "Die Terminalführung erlaubt sowohl einen schnellen Standardlauf als auch eigene Parameter, selektive Kennzahlen, Diagramme und einen PDF-Bericht. Der erfolgreiche Referenzlauf und die 20 bestandenen Unit-Tests zeigen, dass die vorhandenen Kernfunktionen reproduzierbar arbeiten. Gleichzeitig macht die Dokumentation die Modellgrenzen sichtbar und nennt konkrete Schritte für mehr Sicherheit, Konfigurierbarkeit und Realitätsnähe."
    )

    document.add_heading("10.1 Technische Grundlage der Dokumentation", level=2)
    add_body(
        document,
        "Grundlage dieses Berichts sind README.md, main.py, alle Python-Module unter src, die Testdateien unter tests, die enthaltene GPS-Datei sowie die beim Prüflauf erzeugte Terminalausgabe. Der Quellcode verweist beim Rollwiderstandskoeffizienten auf Tengattini und Bigazzi (2018), DOI 10.1080/02640414.2018.1458587. Wetterdaten werden technisch über Open-Meteo und Ortsdaten über Geoapify bezogen."
    )


def build_document() -> Path:
    FIGURES_DIR.mkdir(parents=True, exist_ok=True)
    create_uml_diagram(UML_FILE)
    create_activity_diagram(ACTIVITY_FILE)

    document = Document()
    configure_styles(document)
    configure_section(document.sections[0])

    core = document.core_properties
    core.title = "Projektdokumentation E-Bike-Simulation"
    core.subject = "Abschlussprojekt Programmieren 1"
    core.author = "[Name ergänzen]"
    core.keywords = "Python, E-Bike, GPS, Simulation, Akku, Rekuperation"
    core.comments = "Aus dem verifizierten Projektstand erstellt."

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    add_cover(document)
    add_toc(document)
    add_document_content(document)

    document.save(OUTPUT_FILE)
    return OUTPUT_FILE


if __name__ == "__main__":
    result = build_document()
    print(result)
